##  pip install python-docx

from docx import Document

doc = Document()

doc.add_heading('Document Title', 0)
doc.add_paragraph('This is a docx file')
doc.save('first.docx')
